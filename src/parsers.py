import io
import os
import re
import pypdf
import docx


def clean_pdf_text(text):
    """
    Clean up PDF font ligature and kerning glitches.
    Preserves single-letter words like 'a', 'A', 'I' cleanly.
    """
    if not text:
        return ""

    # Remove null bytes and unprintable control characters
    text = text.replace('\x00', '')
    text = re.sub(r'[\x01-\x08\x0b\x0c\x0e-\x1f]', '', text)

    # Fix specific recurring PDF OCR / kerning glitches
    glitches = {
        "r ainw ater": "rainwater",
        "rainw ater": "rainwater",
        "collec tion": "collection",
        "s ystem": "system",
        "t ypically": "typically",
        "suppor ts": "supports",
        "sanitar y": "sanitary",
        "inspec tion": "inspection",
        "r un-of f": "run-off",
        "coef ficient": "coefficient",
        "over flow": "overflow",
        "ver min": "vermin",
        "ventil ation": "ventilation",
        "ac tivit y": "activity",
        "oper ation": "operation",
        "distr ic t": "district",
        "prov ince": "province",
        "v ill age": "village",
        "par ish": "parish",
        "descr ibe": "describe",
        "af fec ted": "affected",
        "r ainfall": "rainfall",
        "gut ter": "gutter",
        "gut ter ing": "guttering"
    }

    for glitch, fix in glitches.items():
        text = re.sub(re.escape(glitch), fix, text, flags=re.IGNORECASE)

    # Clean up footnote superscript markers like 'source.a' -> 'source.'
    text = re.sub(r'([a-z0-9])\.([a-d1-9])\b', r'\1.', text)

    # Clean excessive whitespace
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\n\s*\n+', '\n\n', text)

    return text.strip()


def fallback_extract_pdf_text(file_bytes_or_path, file_name="document.pdf"):
    """Fallback text extraction for malformed or corrupted PDFs using raw stream parsing."""
    if isinstance(file_bytes_or_path, bytes):
        raw_bytes = file_bytes_or_path
    else:
        try:
            with open(file_bytes_or_path, "rb") as f:
                raw_bytes = f.read()
        except Exception:
            return []

    # Find text sequences inside parentheses or stream blocks
    matches = re.findall(rb'\(([^()\\\r\n]{3,})\)', raw_bytes)
    extracted_strings = []
    for match in matches:
        try:
            s = match.decode('utf-8', errors='ignore').strip()
            if len(s) > 3 and any(c.isalnum() for c in s):
                extracted_strings.append(s)
        except Exception:
            continue

    full_text = " ".join(extracted_strings)
    full_text = clean_pdf_text(full_text)

    if not full_text:
        return []

    # Split into chunks of ~1500 chars to emulate pages
    chunk_size = 1500
    documents = []
    paragraphs = full_text.split(" ")
    current_chunk = []
    current_length = 0
    page_num = 1

    for word in paragraphs:
        current_chunk.append(word)
        current_length += len(word) + 1
        if current_length >= chunk_size:
            chunk_text = " ".join(current_chunk).strip()
            if chunk_text:
                documents.append({
                    "text": chunk_text,
                    "metadata": {
                        "source": file_name,
                        "page": page_num
                    }
                })
                page_num += 1
            current_chunk = []
            current_length = 0

    if current_chunk:
        chunk_text = " ".join(current_chunk).strip()
        if chunk_text:
            documents.append({
                "text": chunk_text,
                "metadata": {
                    "source": file_name,
                    "page": page_num
                }
            })

    return documents


def parse_pdf(file_bytes_or_path, file_name="document.pdf"):
    """Extract text from PDF pages with clean text repair and non-strict error recovery."""
    documents = []

    try:
        if isinstance(file_bytes_or_path, bytes):
            reader = pypdf.PdfReader(io.BytesIO(file_bytes_or_path), strict=False)
        else:
            reader = pypdf.PdfReader(file_bytes_or_path, strict=False)

        try:
            num_pages = len(reader.pages)
        except Exception:
            num_pages = 0

        for i in range(num_pages):
            try:
                page = reader.pages[i]
                text = page.extract_text() or ""
            except Exception:
                text = ""

            text = clean_pdf_text(text)

            if text:
                documents.append({
                    "text": text,
                    "metadata": {
                        "source": file_name,
                        "page": i + 1
                    }
                })
    except Exception as e:
        print(f"[parse_pdf] pypdf failed to parse '{file_name}': {e}")

    # Fallback to stream extraction if pypdf yields no valid text or encounters fatal syntax errors
    if not documents:
        documents = fallback_extract_pdf_text(file_bytes_or_path, file_name=file_name)

    return documents


def parse_docx(file_bytes_or_path, file_name="document.docx"):
    """Extract text from Word DOCX document."""
    documents = []

    if isinstance(file_bytes_or_path, bytes):
        doc = docx.Document(io.BytesIO(file_bytes_or_path))
    else:
        doc = docx.Document(file_bytes_or_path)

    full_text = []
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text.strip())

    text = "\n".join(full_text)
    if text:
        documents.append({
            "text": text,
            "metadata": {
                "source": file_name,
                "page": 1
            }
        })

    return documents


def parse_txt(file_bytes_or_path, file_name="document.txt"):
    """Extract text from plain text or Markdown document."""
    documents = []

    if isinstance(file_bytes_or_path, bytes):
        text = file_bytes_or_path.decode("utf-8", errors="ignore")
    else:
        with open(file_bytes_or_path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()

    text = text.strip()
    if text:
        documents.append({
            "text": text,
            "metadata": {
                "source": file_name,
                "page": 1
            }
        })

    return documents


def parse_document(file_name, file_bytes_or_path):
    """Unified document parser for PDF, DOCX, TXT, and MD files."""
    ext = os.path.splitext(file_name)[1].lower()

    if ext == ".pdf":
        return parse_pdf(file_bytes_or_path, file_name=file_name)
    elif ext in [".docx", ".doc"]:
        return parse_docx(file_bytes_or_path, file_name=file_name)
    elif ext in [".txt", ".md", ".csv", ".json", ".log"]:
        return parse_txt(file_bytes_or_path, file_name=file_name)
    else:
        return parse_txt(file_bytes_or_path, file_name=file_name)
